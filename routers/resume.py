from fastapi import APIRouter, HTTPException, Depends
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional, List
import os
import openai
import models
import auth
import database
from fastapi.responses import FileResponse
from services.job_scraper import SKILL_KEYWORDS
try:
    from docx import Document
    from docx.shared import Pt
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except Exception:
    Document = None
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

router = APIRouter(
    prefix="/resume",
    tags=["Resume"]
)

# ─── Schemas ──────────────────────────────────────────────────

class ResumeRequest(BaseModel):
    full_name: str
    email: str
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    job_title: str
    summary: Optional[str] = None
    experience: List[dict]   # [{title, company, start, end, description}]
    education: List[dict]    # [{degree, institution, year}]
    skills: List[str]
    certifications: Optional[List[str]] = []
    target_job: Optional[str] = None  # Used to tailor resume


class ResumeResponse(BaseModel):
    resume_text: str
    summary_generated: str


class TailoredResumeRequest(BaseModel):
    job_id: Optional[int] = None
    job_title: Optional[str] = None
    job_description: Optional[str] = None
    job_company: Optional[str] = None


# ─── Tailored resume from DB candidate profile ─────────────────

@router.post("/generate-for-job", response_model=ResumeResponse)
async def generate_tailored_resume(
    req: TailoredResumeRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(["candidate"]))
):
    """
    Auto-generate a tailored resume for the logged-in candidate
    based on their DB profile and the target job description.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="OpenAI API key is not configured.")

    # Load candidate profile from DB
    candidate = db.query(models.Candidate).filter(
        models.Candidate.user_id == current_user.id
    ).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate profile not found. Please complete your profile first.")

    if not (candidate.first_name or candidate.last_name):
        raise HTTPException(status_code=400, detail="Please fill in your name in your profile first.")

    # Load job details if job_id provided
    job_title = req.job_title or "Software Professional"
    job_description = req.job_description or ""
    job_company = req.job_company or ""

    if req.job_id:
        job = db.query(models.Job).filter(models.Job.id == req.job_id).first()
        if job:
            job_title = job.title
            job_description = job.description or ""
            job_company = job.company or ""

    # Build candidate details
    full_name = f"{candidate.first_name or ''} {candidate.last_name or ''}".strip()
    skills = candidate.skills if isinstance(candidate.skills, list) else []
    skills_text = ", ".join(skills) if skills else "Not specified"

    # Build the prompt
    target_context = ""
    if job_title or job_description:
        target_context = f"""
TARGET JOB:
Title: {job_title}
Company: {job_company}
Job Description (tailor the resume to match this):
{job_description[:1500]}
"""

    prompt = f"""You are an expert ATS-optimized professional resume writer with 15+ years of experience.

Write a polished, ATS-friendly, professional resume for the candidate below.
{target_context}

CANDIDATE PROFILE (from our database):
Full Name: {full_name}
Email: {current_user.email}
Phone: {candidate.phone or 'Not provided'}
Location: {candidate.location or 'Not provided'}
Current/Desired Title: {candidate.current_title or job_title}
Years of Experience: {candidate.experience_years or 'Not specified'}
Skills: {skills_text}

INSTRUCTIONS:
1. Write the resume in clean professional plain text format with clear section headers (use === or --- dividers).
2. Generate a compelling 3-4 sentence PROFESSIONAL SUMMARY tailored to the job description above.
3. Highlight the most relevant skills from the candidate's profile that match the job description.
4. Structure: CONTACT INFO → PROFESSIONAL SUMMARY → SKILLS → WORK EXPERIENCE → EDUCATION.
5. Since we don't have full work history, create a realistic work experience section based on their title and skill level.
6. Make it ATS-friendly (no tables, no graphics, clean text only).
7. Return ONLY the resume text, no explanations.
"""

    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert professional resume writer who creates highly targeted, ATS-optimized resumes."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        resume_text = response.choices[0].message.content.strip()
    except openai.AuthenticationError:
        raise HTTPException(status_code=401, detail="Invalid OpenAI API key.")
    except openai.RateLimitError:
        raise HTTPException(status_code=429, detail="OpenAI rate limit exceeded. Please try again later.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Resume generation error: {str(e)}")

    # Extract summary
    lines = resume_text.split("\n")
    summary_lines = []
    in_summary = False
    for line in lines:
        if "summary" in line.lower() or "profile" in line.lower():
            in_summary = True
            continue
        if in_summary and line.strip() == "":
            if summary_lines:
                break
            continue
        if in_summary and line.strip().startswith(("SKILLS", "WORK EXP", "EXPERIENCE", "EDUCATION", "===", "---")):
            break
        if in_summary and line.strip():
            summary_lines.append(line.strip())

    summary_generated = " ".join(summary_lines) if summary_lines else "Resume generated successfully."

    return ResumeResponse(
        resume_text=resume_text,
        summary_generated=summary_generated
    )

@router.post("/generate-docx-for-job")
def generate_docx_for_job(
    req: TailoredResumeRequest,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(["candidate"]))
):
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx not available")

    candidate = db.query(models.Candidate).filter(models.Candidate.user_id == current_user.id).first()
    if not candidate:
        raise HTTPException(404, "Candidate profile not found")

    job = None
    if req.job_id:
        job = db.query(models.Job).filter(models.Job.id == req.job_id).first()

    full_name = f"{candidate.first_name or ''} {candidate.last_name or ''}".strip() or "Candidate"
    user = db.query(models.User).filter(models.User.id == candidate.user_id).first()
    email = user.email if user else ""
    phone = candidate.phone or ""
    location = candidate.location or ""
    title = candidate.current_title or (req.job_title or "Professional")
    company = (job.company if job else req.job_company) or ""
    job_title = (job.title if job else req.job_title) or title
    job_date = "Present"
    skills = candidate.skills if isinstance(candidate.skills, list) else []
    summary = f"{title} with experience delivering results. Skilled in {', '.join(skills)}. Tailored for {job_title} at {company}."
    job_description = (job.description if job else req.job_description) or ""
    job_skills = []
    if job and job.skills:
        if isinstance(job.skills, list):
            job_skills = [str(s) for s in job.skills]
    else:
        text = job_description.lower()
        job_skills = [k.title() for k in SKILL_KEYWORDS if k in text]
    keywords = list(dict.fromkeys((skills or []) + job_skills))

    parsed_summary = None
    parsed_experience = []
    parsed_education = None
    resume_path = None
    if candidate.resume_url and candidate.resume_url.startswith("/static/"):
        p = os.path.join("uploads", candidate.resume_url.replace("/static/", ""))
        if os.path.exists(p):
            resume_path = p
            ext = p.lower().split(".")[-1]
            content = ""
            if ext == "txt":
                try:
                    with open(p, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                except Exception:
                    content = ""
            elif ext == "docx" and Document is not None:
                try:
                    d = Document(p)
                    content = "\n".join([para.text for para in d.paragraphs])
                except Exception:
                    content = ""
            elif ext == "pdf" and pdf_extract_text is not None:
                try:
                    content = pdf_extract_text(p) or ""
                except Exception:
                    content = ""
            text = content.strip()
            if text:
                lower = text.lower()
                def extract_section(name_list):
                    for name in name_list:
                        i = lower.find(name)
                        if i != -1:
                            return i
                    return -1
                idx_summary = extract_section(["professional summary", "summary"])
                idx_skills = extract_section(["technical skills", "skills"])
                idx_exp = extract_section(["work experience", "experience"])
                idx_edu = extract_section(["education"])
                if idx_summary >= 0:
                    end = idx_skills if idx_skills > idx_summary else (idx_exp if idx_exp > idx_summary else len(text))
                    parsed_summary = text[idx_summary:end].split("\n", 1)[-1].strip()[:600]
                if idx_edu >= 0:
                    end = len(text)
                    parsed_education = text[idx_edu:end].strip()[:400]
                if idx_exp >= 0:
                    end = idx_edu if idx_edu > idx_exp else len(text)
                    exp_text = text[idx_exp:end]
                    lines = [l.strip() for l in exp_text.split("\n") if l.strip()]
                    bullets = []
                    for l in lines:
                        if l.startswith(("•", "-", "*")) or l[:2].lower() in ["- ", "* "]:
                            bullets.append(l.lstrip("•-* ").strip())
                    if not bullets:
                        bullets = lines
                    parsed_experience = bullets[:14]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    doc.styles['Normal'].font.size = Pt(10)
    h = doc.add_heading(full_name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"{email} | {phone} | {location}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(parsed_summary or summary)
    doc.add_heading("Technical Skills", level=1)
    table = doc.add_table(rows=max(1, (len(skills)+1)//2), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r = 0; c = 0
    for s in skills:
        cell = table.cell(r, c)
        cell.text = s
        c = (c + 1) % 2
        if c == 0:
            r += 1

    doc.add_heading("Job Keywords", level=1)
    for kw in keywords[:12]:
        doc.add_paragraph(kw, style='List Bullet')

    doc.add_heading("Work Experience", level=1)
    exp_title = doc.add_paragraph()
    run = exp_title.add_run(f"{job_title}")
    run.bold = True
    exp_title.add_run(f"    {job_date}").italic = True
    doc.add_paragraph(company).italic = True
    default_bullets = [
        "Analyzed business and data workflows aligned to role requirements.",
        "Mapped candidate experience to job scope and deliverables.",
        "Built dashboards and reports using relevant BI tools.",
        "Collaborated across teams to deliver outcomes.",
    ]
    bullets = parsed_experience if parsed_experience else default_bullets
    for bullet in bullets[:12]:
        doc.add_paragraph(bullet, style='List Bullet')

    if job_description:
        doc.add_heading("Job Description Highlights", level=1)
        doc.add_paragraph(job_description[:900])

    doc.add_heading("Education", level=1)
    doc.add_paragraph(parsed_education or "M.S. or Bachelor’s degree – Details")

    out_dir = os.path.join("uploads", "resumes_generated")
    os.makedirs(out_dir, exist_ok=True)
    fname = f"{candidate.id}_{job.id if job else 'job'}_{job_title.replace(' ', '_')}.docx"
    out_path = os.path.join(out_dir, fname)
    doc.save(out_path)
    return FileResponse(out_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"Resume_{job_title.replace(' ', '_')}.docx")

@router.post("/generate-batch-for-candidates")
async def generate_batch_resumes_for_candidates(
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.require_role(["admin"]))
):
    """
    (Admin Only) Iterates through all candidates, finds their top job matches, 
    and generates tailored resumes. Returns a status mapping.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="OpenAI API key is not configured.")

    client = openai.OpenAI(api_key=api_key)
    
    candidates = db.query(models.Candidate).all()
    results = []
    
    from services.job_scraper import get_matched_jobs_for_candidate
    
    for candidate in candidates:
        if not candidate.skills or not candidate.first_name:
            continue
            
        # Get top 3 matched jobs for this candidate
        matched_jobs = get_matched_jobs_for_candidate(db, candidate, min_score=20, top_n=3)
        if not matched_jobs:
            continue
            
        full_name = f"{candidate.first_name or ''} {candidate.last_name or ''}".strip()
        skills = candidate.skills if isinstance(candidate.skills, list) else []
        skills_text = ", ".join(skills) if skills else "Not specified"
        
        candidate_resumes = []
        
        for job_data in matched_jobs:
            target_context = f"""
TARGET JOB:
Title: {job_data['title']}
Company: {job_data.get('company', '')}
Job Description (tailor the resume to match this):
{job_data.get('description', '')[:1000]}
"""

            prompt = f"""You are an expert ATS-optimized professional resume writer with 15+ years of experience.

Write a polished, ATS-friendly, professional resume for the candidate below.
{target_context}

CANDIDATE PROFILE (from our database):
Full Name: {full_name}
Phone: {candidate.phone or 'Not provided'}
Location: {candidate.location or 'Not provided'}
Current/Desired Title: {candidate.current_title or job_data['title']}
Years of Experience: {candidate.experience_years or 'Not specified'}
Skills: {skills_text}

INSTRUCTIONS:
1. Write the resume in clean professional plain text format with clear section headers.
2. Generate a compelling 3-4 sentence PROFESSIONAL SUMMARY tailored to the job description above.
3. Highlight the most relevant skills from the candidate's profile that match the job description.
4. Structure: CONTACT INFO → PROFESSIONAL SUMMARY → SKILLS → WORK EXPERIENCE → EDUCATION.
5. Create a realistic work experience section based on their title and skill level.
6. Make it ATS-friendly (no tables, no graphics, clean text only).
7. Return ONLY the resume text.
"""
            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini", # use mini for batch to save time/cost
                    messages=[
                        {"role": "system", "content": "You are an expert professional resume writer."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=1500
                )
                resume_text = response.choices[0].message.content.strip()
                
                candidate_resumes.append({
                    "job_id": job_data["job_id"],
                    "job_title": job_data["title"],
                    "company": job_data["company"],
                    "resume_excerpt": resume_text[:200] + "..."
                })
            except Exception as e:
                candidate_resumes.append({
                    "job_id": job_data["job_id"],
                    "error": str(e)
                })

        results.append({
            "candidate_id": candidate.id,
            "candidate_name": full_name,
            "resumes_generated": candidate_resumes
        })

    return {
        "message": f"Batch processing complete. Processed {len(results)} candidates.",
        "details": results
    }


# ─── Original manual resume generation ────────────────────────

@router.post("/generate", response_model=ResumeResponse)
async def generate_resume(req: ResumeRequest):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="OpenAI API key is not configured.")

    client = openai.OpenAI(api_key=api_key)

    exp_text = "\n".join([
        f"- {e.get('title','N/A')} at {e.get('company','N/A')} ({e.get('start','')}-{e.get('end','Present')}): {e.get('description','')}"
        for e in req.experience
    ])
    edu_text = "\n".join([
        f"- {e.get('degree','N/A')} from {e.get('institution','N/A')} ({e.get('year','')})"
        for e in req.education
    ])
    skills_text = ", ".join(req.skills)
    certs_text = ", ".join(req.certifications) if req.certifications else "None"

    target_context = f"\nThe resume should be tailored for the role: **{req.target_job}**." if req.target_job else ""

    prompt = f"""You are an expert professional resume writer with 15+ years of experience helping candidates land top-tier jobs.

Write a polished, ATS-friendly resume for the following person.{target_context}

CANDIDATE INFORMATION:
Name: {req.full_name}
Email: {req.email}
Phone: {req.phone or 'N/A'}
Location: {req.location or 'N/A'}
LinkedIn: {req.linkedin or 'N/A'}
Current/Desired Title: {req.job_title}
User Summary: {req.summary or 'Generate a strong professional summary based on their background.'}

WORK EXPERIENCE:
{exp_text}

EDUCATION:
{edu_text}

SKILLS:
{skills_text}

CERTIFICATIONS:
{certs_text}

INSTRUCTIONS:
1. Write the resume in a clean, professional format using plain text with clear section headers.
2. Use strong action verbs and quantify achievements where possible.
3. Generate a compelling 3-4 sentence professional summary at the top.
4. Structure: Contact Info → Professional Summary → Skills → Work Experience → Education → Certifications.
5. Make it suitable for ATS (avoid tables, graphics).
6. Return ONLY the resume text, nothing else.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert professional resume writer."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        resume_text = response.choices[0].message.content.strip()
    except openai.AuthenticationError:
        raise HTTPException(status_code=401, detail="Invalid OpenAI API key.")
    except openai.RateLimitError:
        raise HTTPException(status_code=429, detail="OpenAI rate limit exceeded. Please try again later.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI API error: {str(e)}")

    lines = resume_text.split("\n")
    summary_lines = []
    in_summary = False
    for line in lines:
        if "summary" in line.lower() or "profile" in line.lower():
            in_summary = True
            continue
        if in_summary and line.strip() == "":
            if summary_lines:
                break
            continue
        if in_summary and line.strip().startswith(("SKILLS", "WORK EXP", "EXPERIENCE", "EDUCATION", "---")):
            break
        if in_summary and line.strip():
            summary_lines.append(line.strip())

    summary_generated = " ".join(summary_lines) if summary_lines else "Resume generated successfully."

    return ResumeResponse(
        resume_text=resume_text,
        summary_generated=summary_generated
    )
